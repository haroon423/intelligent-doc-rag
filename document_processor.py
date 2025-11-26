import os
import uuid
from typing import List, Dict, Any, Optional
import pypdf
import docx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

class DocumentProcessor:
    """Process various document formats and prepare them for vector storage"""
    
    def __init__(self, chunk_size=1000, chunk_overlap=200):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
    
    def load_pdf(self, file_path: str) -> List[Document]:
        """Extract text from PDF file"""
        documents = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
                # Create metadata
                metadata = {
                    "source": file_path,
                    "file_type": "pdf",
                    "page_count": len(pdf_reader.pages)
                }
                
                documents.append(Document(page_content=text, metadata=metadata))
        except Exception as e:
            print(f"Error processing PDF {file_path}: {str(e)}")
        
        return documents
    
    def load_docx(self, file_path: str) -> List[Document]:
        """Extract text from DOCX file"""
        documents = []
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Create metadata
            metadata = {
                "source": file_path,
                "file_type": "docx"
            }
            
            documents.append(Document(page_content=text, metadata=metadata))
        except Exception as e:
            print(f"Error processing DOCX {file_path}: {str(e)}")
        
        return documents
    
    def load_txt(self, file_path: str) -> List[Document]:
        """Extract text from TXT file"""
        documents = []
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            # Create metadata
            metadata = {
                "source": file_path,
                "file_type": "txt"
            }
            
            documents.append(Document(page_content=text, metadata=metadata))
        except Exception as e:
            print(f"Error processing TXT {file_path}: {str(e)}")
        
        return documents
    
    def process_documents(self, file_paths: List[str]) -> List[Document]:
        """Process multiple documents and split into chunks"""
        all_documents = []
        
        for file_path in file_paths:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                documents = self.load_pdf(file_path)
            elif file_ext == '.docx':
                documents = self.load_docx(file_path)
            elif file_ext == '.txt':
                documents = self.load_txt(file_path)
            else:
                print(f"Unsupported file type: {file_ext}")
                continue
            
            all_documents.extend(documents)
        
        # Split documents into chunks
        chunks = self.text_splitter.split_documents(all_documents)
        
        # Add unique IDs to chunks
        for i, chunk in enumerate(chunks):
            chunk.metadata["chunk_id"] = str(uuid.uuid4())
            chunk.metadata["chunk_index"] = i
        
        return chunks
    
    def create_document_from_text(self, text: str, source: str = "manual_input") -> List[Document]:
        """Create document chunks from raw text"""
        metadata = {
            "source": source,
            "file_type": "text"
        }
        
        document = Document(page_content=text, metadata=metadata)
        chunks = self.text_splitter.split_documents([document])
        
        # Add unique IDs to chunks
        for i, chunk in enumerate(chunks):
            chunk.metadata["chunk_id"] = str(uuid.uuid4())
            chunk.metadata["chunk_index"] = i
        
        return chunks